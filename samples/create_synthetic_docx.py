"""Create a synthetic DOCX for local Blue Linter testing."""

from pathlib import Path

from docx import Document

OUTPUT_PATH = Path(__file__).with_name("synthetic-style-review.docx")


def main() -> None:
    document = Document()

    document.add_heading("Quarterly business REVIEW", level=1)
    document.add_paragraph("Revenue increased by 25 %  while margin improved by 3 %.")
    document.add_paragraph("The API response time improved, but API ownership remains unclear.")
    document.add_paragraph("")

    document.add_heading("Operating Model", level=2)
    document.add_paragraph("Reduce operating costs", style="List Bullet")
    document.add_paragraph("Improve customer retention.", style="List Bullet")
    document.add_paragraph("1. Manual numbered item without punctuation")
    document.add_paragraph("A) Manual lettered item.")

    document.add_heading("Defined Terms", level=2)
    document.add_paragraph("Application Programming Interface (API) guidance is available.")
    document.add_paragraph("The API glossary entry should not create a new acronym finding.")

    document.save(OUTPUT_PATH)
    print(f"Created {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
