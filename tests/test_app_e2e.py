import json
from pathlib import Path
from zipfile import ZipFile

from docx import Document

from blue_linter.app import run_review
from blue_linter.parser import parse_docx

EXPECTED_ARCHIVE_PATHS = [
    "original/document-original.docx",
    "candidate/document-style-corrected-candidate.docx",
    "reports/style-analysis-report.html",
    "reports/style-analysis-report.json",
    "reports/validation-report.html",
    "rules/active-style-rules.yaml",
    "audit/audit.json",
]


def save_synthetic_document(path: Path) -> Path:
    document = Document()
    document.add_heading("Quarterly business REVIEW", level=1)
    document.add_paragraph("Revenue increased by 25 %  while margin improved.")
    document.add_paragraph("The API remains undocumented.")
    document.save(path)
    return path


def test_run_review_creates_complete_review_package(tmp_path: Path) -> None:
    input_path = save_synthetic_document(tmp_path / "input.docx")
    output_path = tmp_path / "style-review-package.zip"

    result = run_review(input_path=input_path, output_path=output_path)

    assert result.status == "completed"
    assert result.output_path == output_path
    assert result.finding_count > 0
    assert result.applied_fix_count > 0
    assert result.validation_remaining_count < result.finding_count
    assert result.packaged_paths == EXPECTED_ARCHIVE_PATHS

    with ZipFile(output_path) as package:
        assert package.namelist() == EXPECTED_ARCHIVE_PATHS
        for archive_path in EXPECTED_ARCHIVE_PATHS:
            assert package.getinfo(archive_path).file_size > 0

        assert package.read("rules/active-style-rules.yaml") == Path(
            "rules/active-style-rules.yaml"
        ).read_bytes()

        candidate_path = tmp_path / "candidate-from-package.docx"
        candidate_path.write_bytes(
            package.read("candidate/document-style-corrected-candidate.docx")
        )
        candidate_text = "\n".join(block.text for block in parse_docx(candidate_path).blocks)
        assert "25%" in candidate_text
        assert "25 %" not in candidate_text

        analysis_payload = json.loads(
            package.read("reports/style-analysis-report.json").decode("utf-8")
        )
        audit_payload = json.loads(package.read("audit/audit.json").decode("utf-8"))
        validation_html = package.read("reports/validation-report.html").decode("utf-8")

    assert analysis_payload["summary"]["total_findings"] == result.finding_count
    assert audit_payload["counts"]["validation"]["total_findings"] == (
        result.validation_remaining_count
    )
    assert "Validation Report" in validation_html
