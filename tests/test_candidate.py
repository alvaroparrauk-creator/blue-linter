from pathlib import Path

from docx import Document

from blue_linter.candidate import generate_candidate_document
from blue_linter.engine import RuleEngine
from blue_linter.parser import parse_docx
from blue_linter.rules import StyleRule


def save_document(path: Path, paragraphs: list[str]) -> Path:
    document = Document()
    for text in paragraphs:
        document.add_paragraph(text)
    document.save(path)
    return path


def save_multi_run_document(path: Path) -> Path:
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("Revenue increased by ")
    paragraph.add_run("25 %")
    document.save(path)
    return path


def percent_rule() -> StyleRule:
    return StyleRule(
        id="STYLE-PERCENT-SPACING",
        version="1.0.0",
        name="Percentage spacing",
        category="number-format",
        severity="medium",
        enabled=True,
        type="regex_replace",
        auto_fix=True,
        pattern=r"\b(\d+)\s+%",
        replacement=r"\1%",
    )


def whitespace_rule() -> StyleRule:
    return StyleRule(
        id="STYLE-REPEATED-WHITESPACE",
        version="1.0.0",
        name="Repeated whitespace",
        category="spacing",
        severity="low",
        enabled=True,
        type="regex_replace",
        auto_fix=True,
        pattern=r" {2,}",
        replacement=" ",
    )


def manual_rule() -> StyleRule:
    return StyleRule(
        id="STYLE-ACRONYM-FIRST-USE",
        version="1.0.0",
        name="Acronym first use",
        category="terminology",
        severity="high",
        enabled=True,
        type="acronym_first_use",
        auto_fix=False,
    )


def run_rules(path: Path, rules: list[StyleRule]):
    parsed_document = parse_docx(path)
    return parsed_document, RuleEngine().run(parsed_document, rules)


def test_candidate_file_is_created_and_original_remains_unchanged(tmp_path: Path) -> None:
    input_path = save_document(tmp_path / "original.docx", ["Revenue increased by 25 %."])
    output_path = tmp_path / "candidate.docx"
    rules = [percent_rule()]
    parsed_document, findings = run_rules(input_path, rules)

    result = generate_candidate_document(input_path, output_path, parsed_document, rules, findings)

    assert result.candidate_path == output_path
    assert output_path.exists()
    assert parse_docx(input_path).blocks[0].text == "Revenue increased by 25 %."
    assert parse_docx(output_path).blocks[0].text == "Revenue increased by 25%."


def test_multiple_auto_fix_rules_on_same_paragraph_combine_in_rule_order(tmp_path: Path) -> None:
    input_path = save_document(
        tmp_path / "original.docx",
        ["Revenue increased by 25 %  while margin improved by 3 %."],
    )
    output_path = tmp_path / "candidate.docx"
    rules = [percent_rule(), whitespace_rule()]
    parsed_document, findings = run_rules(input_path, rules)

    result = generate_candidate_document(input_path, output_path, parsed_document, rules, findings)

    assert result.applied_count == 2
    assert result.skipped_count == 0
    assert parse_docx(output_path).blocks[0].text == (
        "Revenue increased by 25% while margin improved by 3%."
    )
    assert [finding.review_status for finding in result.findings] == [
        "auto_applied",
        "auto_applied",
    ]


def test_manual_review_findings_are_not_applied(tmp_path: Path) -> None:
    input_path = save_document(tmp_path / "original.docx", ["The API remains undocumented."])
    output_path = tmp_path / "candidate.docx"
    rules = [manual_rule()]
    parsed_document, findings = run_rules(input_path, rules)

    result = generate_candidate_document(input_path, output_path, parsed_document, rules, findings)

    assert result.applied_count == 0
    assert result.skipped_count == 0
    assert parse_docx(output_path).blocks[0].text == "The API remains undocumented."
    assert result.findings[0].review_status == "manual_review_required"


def test_multi_run_paragraph_is_skipped_to_preserve_formatting(tmp_path: Path) -> None:
    input_path = save_multi_run_document(tmp_path / "original.docx")
    output_path = tmp_path / "candidate.docx"
    rules = [percent_rule()]
    parsed_document, findings = run_rules(input_path, rules)

    result = generate_candidate_document(input_path, output_path, parsed_document, rules, findings)

    assert result.applied_count == 0
    assert result.skipped_count == 1
    assert result.skipped_reasons == {
        "F-000001": "Paragraph has multiple runs; formatting preserved.",
    }
    assert parse_docx(output_path).blocks[0].text == "Revenue increased by 25 %"


def test_stale_paragraph_location_is_skipped(tmp_path: Path) -> None:
    input_path = save_document(tmp_path / "original.docx", ["Revenue increased by 25 %."])
    output_path = tmp_path / "candidate.docx"
    rules = [percent_rule()]
    parsed_document, findings = run_rules(input_path, rules)
    stale_finding = findings[0].model_copy(deep=True)
    stale_finding.location.paragraph_index = 7

    result = generate_candidate_document(
        input_path,
        output_path,
        parsed_document,
        rules,
        [stale_finding],
    )

    assert result.applied_count == 0
    assert result.skipped_reasons == {
        "F-000001": "Finding paragraph index does not match parsed block.",
    }


def test_mismatched_paragraph_text_is_skipped(tmp_path: Path) -> None:
    input_path = save_document(tmp_path / "original.docx", ["Revenue increased by 25 %."])
    output_path = tmp_path / "candidate.docx"
    rules = [percent_rule()]
    parsed_document, findings = run_rules(input_path, rules)
    parsed_document.blocks[0].text = "Different parsed text"

    result = generate_candidate_document(input_path, output_path, parsed_document, rules, findings)

    assert result.applied_count == 0
    assert result.skipped_reasons == {
        "F-000001": "Paragraph text no longer matches parsed text.",
    }


def test_original_findings_are_not_mutated(tmp_path: Path) -> None:
    input_path = save_document(tmp_path / "original.docx", ["Revenue increased by 25 %."])
    output_path = tmp_path / "candidate.docx"
    rules = [percent_rule()]
    parsed_document, findings = run_rules(input_path, rules)

    result = generate_candidate_document(input_path, output_path, parsed_document, rules, findings)

    assert findings[0].applied_to_candidate is False
    assert findings[0].review_status == "pending_review"
    assert result.findings[0].applied_to_candidate is True
    assert result.findings[0].review_status == "auto_applied"
