"""DOCX parsing helpers for Blue Linter."""

import re
from pathlib import Path
from zipfile import BadZipFile

from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from docx.text.paragraph import Paragraph

from blue_linter.document import BlockType, DocumentBlock, ParsedDocument

BULLET_PREFIXES = ("\u2022", "-", "*", "\u2013")
LIST_PREFIX_PATTERN = re.compile(r"^(?:\d+|[A-Za-z])[\.)]")


class DocumentParseError(ValueError):
    """Raised when a DOCX file cannot be parsed."""


def parse_docx(path: Path) -> ParsedDocument:
    """Parse main-body DOCX paragraphs into Blue Linter's internal document model."""
    if path.suffix.lower() != ".docx":
        raise DocumentParseError("Input document must use the .docx extension.")
    if not path.exists():
        raise DocumentParseError(f"Input document does not exist: {path}")
    if not path.is_file():
        raise DocumentParseError(f"Input path is not a file: {path}")

    try:
        document = Document(path)
    except (OSError, PackageNotFoundError, BadZipFile, KeyError) as exc:
        raise DocumentParseError(f"Unable to parse DOCX document at {path}: {exc}") from exc

    blocks: list[DocumentBlock] = []
    section_title: str | None = None
    for paragraph_index, paragraph in enumerate(document.paragraphs):
        text = paragraph.text
        style_name = _style_name(paragraph)
        block_type = _classify_block(paragraph, text, style_name)

        if block_type == "heading" and text.strip():
            section_title = text

        blocks.append(
            DocumentBlock(
                block_id=f"p-{paragraph_index + 1:06d}",
                paragraph_index=paragraph_index,
                text=text,
                style_name=style_name,
                block_type=block_type,
                section_title=section_title,
            )
        )

    return ParsedDocument(source_name=path.name, blocks=blocks)


def _style_name(paragraph: Paragraph) -> str:
    try:
        style = paragraph.style
        name = style.name if style is not None else None
    except (AttributeError, KeyError):
        name = None
    return name or "Unknown"


def _classify_block(paragraph: Paragraph, text: str, style_name: str) -> BlockType:
    stripped_text = text.strip()
    normalized_style = style_name.lower()

    if not stripped_text:
        return "unknown"
    if normalized_style.startswith("heading"):
        return "heading"
    if _is_bullet_style(normalized_style) or _has_bullet_prefix(stripped_text):
        return "bullet"
    if (
        _has_numbering_metadata(paragraph)
        or _is_list_style(normalized_style)
        or _has_list_prefix(stripped_text)
    ):
        return "list_item"
    return "paragraph"


def _is_bullet_style(normalized_style: str) -> bool:
    return "bullet" in normalized_style


def _is_list_style(normalized_style: str) -> bool:
    return "list" in normalized_style or "number" in normalized_style


def _has_bullet_prefix(stripped_text: str) -> bool:
    return stripped_text.startswith(BULLET_PREFIXES)


def _has_list_prefix(stripped_text: str) -> bool:
    return LIST_PREFIX_PATTERN.match(stripped_text) is not None


def _has_numbering_metadata(paragraph: Paragraph) -> bool:
    paragraph_properties = paragraph._p.pPr  # pyright: ignore[reportPrivateUsage]
    return paragraph_properties is not None and paragraph_properties.numPr is not None
