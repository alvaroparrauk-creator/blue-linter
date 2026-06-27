"""Candidate DOCX generation for safe deterministic fixes."""

from __future__ import annotations

import re
import shutil
from collections import defaultdict
from pathlib import Path
from zipfile import BadZipFile

from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from pydantic import BaseModel, ConfigDict, Field

from blue_linter.document import DocumentBlock, ParsedDocument
from blue_linter.findings import Finding
from blue_linter.rules import StyleRule


class CandidateGenerationError(ValueError):
    """Raised when a candidate DOCX cannot be generated."""


class CandidateGenerationResult(BaseModel):
    """Result returned after creating a candidate document."""

    model_config = ConfigDict(extra="forbid")

    candidate_path: Path
    findings: list[Finding]
    applied_count: int = Field(ge=0)
    skipped_count: int = Field(ge=0)
    skipped_reasons: dict[str, str]


def generate_candidate_document(
    input_path: Path,
    output_path: Path,
    parsed_document: ParsedDocument,
    active_rules: list[StyleRule],
    findings: list[Finding],
) -> CandidateGenerationResult:
    """Create a candidate DOCX by applying safe auto-fix findings."""
    _validate_docx_path(input_path, label="Input")
    if output_path.suffix.lower() != ".docx":
        raise CandidateGenerationError("Output candidate document must use the .docx extension.")

    try:
        output_path.parent.mkdir(parents=True, exist_ok=True)
        shutil.copyfile(input_path, output_path)
        document = Document(output_path)
    except (OSError, PackageNotFoundError, BadZipFile, KeyError) as exc:
        raise CandidateGenerationError(f"Unable to prepare candidate document: {exc}") from exc

    block_by_id = {block.block_id: block for block in parsed_document.blocks}
    auto_fix_rules = _auto_fix_rules_by_id(active_rules)
    candidate_findings = [finding.model_copy(deep=True) for finding in findings]
    finding_by_id = {finding.finding_id: finding for finding in candidate_findings}
    findings_by_block = _eligible_findings_by_block(candidate_findings, auto_fix_rules)

    skipped_reasons: dict[str, str] = {}
    applied_finding_ids: set[str] = set()

    for block_id, block_findings in findings_by_block.items():
        block = block_by_id.get(block_id)
        if block is None:
            _skip_all(block_findings, skipped_reasons, "Parsed block is no longer available.")
            continue
        stale_paragraph_index = any(
            finding.location.paragraph_index != block.paragraph_index
            for finding in block_findings
        )
        if stale_paragraph_index:
            _skip_all(
                block_findings,
                skipped_reasons,
                "Finding paragraph index does not match parsed block.",
            )
            continue
        if block.paragraph_index >= len(document.paragraphs):
            _skip_all(
                block_findings,
                skipped_reasons,
                "Paragraph index is outside the candidate document.",
            )
            continue

        paragraph = document.paragraphs[block.paragraph_index]
        if paragraph.text != block.text:
            _skip_all(
                block_findings,
                skipped_reasons,
                "Paragraph text no longer matches parsed text.",
            )
            continue
        if len(paragraph.runs) > 1:
            _skip_all(
                block_findings,
                skipped_reasons,
                "Paragraph has multiple runs; formatting preserved.",
            )
            continue

        try:
            candidate_text = _apply_rules_in_order(block, block_findings, active_rules)
        except re.error as exc:
            raise CandidateGenerationError(f"Invalid auto-fix regex pattern: {exc}") from exc
        if candidate_text == block.text:
            _skip_all(block_findings, skipped_reasons, "Auto-fix rules produced no text change.")
            continue

        if len(paragraph.runs) == 1:
            paragraph.runs[0].text = candidate_text
        else:
            paragraph.add_run(candidate_text)

        for finding in block_findings:
            applied_finding_ids.add(finding.finding_id)

    for finding_id in applied_finding_ids:
        finding = finding_by_id[finding_id]
        finding.applied_to_candidate = True
        finding.review_status = "auto_applied"

    try:
        document.save(output_path)
    except OSError as exc:
        raise CandidateGenerationError(
            f"Unable to write candidate document at {output_path}: {exc}"
        ) from exc

    return CandidateGenerationResult(
        candidate_path=output_path,
        findings=candidate_findings,
        applied_count=len(applied_finding_ids),
        skipped_count=len(skipped_reasons),
        skipped_reasons=skipped_reasons,
    )


def _validate_docx_path(path: Path, *, label: str) -> None:
    if path.suffix.lower() != ".docx":
        raise CandidateGenerationError(f"{label} document must use the .docx extension.")
    if not path.exists():
        raise CandidateGenerationError(f"{label} document does not exist: {path}")
    if not path.is_file():
        raise CandidateGenerationError(f"{label} path is not a file: {path}")


def _auto_fix_rules_by_id(active_rules: list[StyleRule]) -> dict[tuple[str, str], StyleRule]:
    return {
        (rule.id, rule.version): rule
        for rule in active_rules
        if rule.enabled
        and rule.auto_fix
        and rule.type == "regex_replace"
        and rule.pattern is not None
        and rule.replacement is not None
    }


def _eligible_findings_by_block(
    findings: list[Finding],
    auto_fix_rules: dict[tuple[str, str], StyleRule],
) -> dict[str, list[Finding]]:
    findings_by_block: dict[str, list[Finding]] = defaultdict(list)
    for finding in findings:
        rule_key = (finding.rule_id, finding.rule_version)
        if rule_key not in auto_fix_rules or finding.suggested_correction is None:
            continue
        findings_by_block[finding.location.block_id].append(finding)
    return findings_by_block


def _apply_rules_in_order(
    block: DocumentBlock,
    block_findings: list[Finding],
    active_rules: list[StyleRule],
) -> str:
    text = block.text
    finding_rule_keys = {
        (finding.rule_id, finding.rule_version)
        for finding in block_findings
    }
    for rule in active_rules:
        rule_key = (rule.id, rule.version)
        if rule_key not in finding_rule_keys:
            continue
        if not rule.auto_fix or rule.type != "regex_replace":
            continue
        if rule.pattern is None or rule.replacement is None:
            continue
        text = re.sub(rule.pattern, rule.replacement, text)
    return text


def _skip_all(
    findings: list[Finding],
    skipped_reasons: dict[str, str],
    reason: str,
) -> None:
    for finding in findings:
        skipped_reasons[finding.finding_id] = reason
